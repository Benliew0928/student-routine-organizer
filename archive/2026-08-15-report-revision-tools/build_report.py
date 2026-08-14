from __future__ import annotations

import copy
import shutil
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt


TNR = "Times New Roman"


def set_rfonts(rpr, name: str = TNR) -> None:
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), name)
    for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        rfonts.attrib.pop(qn(f"w:{attr}"), None)


def set_run_font(run, name: str = TNR, size: Pt | None = None) -> None:
    run.font.name = name
    if size is not None:
        run.font.size = size
    rpr = run._element.get_or_add_rPr()
    set_rfonts(rpr, name)


def clone_with_plain_text(element, text: str):
    cloned = copy.deepcopy(element)
    ppr = cloned.find(qn("w:pPr"))
    for child in list(cloned):
        if child is not ppr:
            cloned.remove(child)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    set_rfonts(rpr)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    cloned.append(run)
    return cloned


def add_before_sectpr(doc: Document, element) -> None:
    doc.element.body.insert_element_before(element, "w:sectPr")


def set_a4(section, landscape: bool) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    if landscape:
        section.page_width = Mm(297)
        section.page_height = Mm(210)
    else:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)


def new_section(doc: Document, landscape: bool):
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    set_a4(sec, landscape)
    sec.header.is_linked_to_previous = True
    sec.footer.is_linked_to_previous = True
    return sec


def add_picture(doc: Document, image_path: Path, width_inches: float = 8.8, page_break_before: bool = False) -> None:
    p = doc.add_paragraph(style="Figure" if "Figure" in [s.name for s in doc.styles] else None)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.page_break_before = page_break_before
    p.add_run().add_picture(str(image_path), width=Inches(width_inches))


def resize_inline_drawing(element, width_inches: float):
    extent = element.find(".//" + qn("wp:extent"))
    if extent is None:
        return element
    old_cx = int(extent.get("cx"))
    old_cy = int(extent.get("cy"))
    new_cx = int(round(width_inches * 914400))
    new_cy = int(round(old_cy * new_cx / old_cx))
    extent.set("cx", str(new_cx))
    extent.set("cy", str(new_cy))
    for aext in element.findall(".//" + qn("a:ext")):
        aext.set("cx", str(new_cx))
        aext.set("cy", str(new_cy))
    return element


def remove_table_shading(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            tcpr = cell._tc.get_or_add_tcPr()
            shd = tcpr.find(qn("w:shd"))
            if shd is not None:
                tcpr.remove(shd)


def set_table_rules(table, report_table: bool) -> None:
    table.autofit = False
    for ridx, row in enumerate(table.rows):
        trpr = row._tr.get_or_add_trPr()
        if trpr.find(qn("w:cantSplit")) is None:
            trpr.append(OxmlElement("w:cantSplit"))
        if report_table and ridx == 0 and trpr.find(qn("w:tblHeader")) is None:
            hdr = OxmlElement("w:tblHeader")
            hdr.set(qn("w:val"), "true")
            trpr.append(hdr)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_run_font(run, size=Pt(12))
                    if report_table and ridx == 0:
                        run.bold = True
    if report_table:
        remove_table_shading(table)


def configure_styles(doc: Document) -> None:
    for style in doc.styles:
        if style.type not in (WD_STYLE_TYPE.PARAGRAPH, WD_STYLE_TYPE.CHARACTER, WD_STYLE_TYPE.TABLE):
            continue
        try:
            style.font.name = TNR
            rpr = style.element.get_or_add_rPr()
            set_rfonts(rpr)
        except Exception:
            pass

    normal = doc.styles["Normal"]
    normal.font.name = TNR
    normal.font.size = Pt(12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)

    for name, size in (("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 12)):
        if name in [s.name for s in doc.styles]:
            st = doc.styles[name]
            st.font.name = TNR
            st.font.size = Pt(size)
            st.font.bold = True
            st.paragraph_format.keep_with_next = True
            st.paragraph_format.keep_together = True
            st.paragraph_format.line_spacing = 1.0
            st.paragraph_format.space_after = Pt(6)
    doc.styles["Heading 1"].paragraph_format.page_break_before = True
    doc.styles["Heading 2"].paragraph_format.page_break_before = False
    doc.styles["Heading 3"].paragraph_format.page_break_before = False

    for name in ("Caption", "Caption 2", "Caption2"):
        if name in [s.name for s in doc.styles]:
            st = doc.styles[name]
            st.font.name = TNR
            st.font.size = Pt(10)
            st.font.italic = True
            st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            st.paragraph_format.line_spacing = 1.0
            st.paragraph_format.space_before = Pt(2)
            st.paragraph_format.space_after = Pt(6)
            st.paragraph_format.keep_together = True

    for name in ("TOC 1", "TOC1", "TOC 2", "TOC2", "TOC 3", "TOC3", "Table of Figures", "TableofFigures"):
        if name in [s.name for s in doc.styles]:
            st = doc.styles[name]
            st.font.name = TNR
            st.font.size = Pt(12)
            st.paragraph_format.line_spacing = 1.0
            st.paragraph_format.space_before = Pt(0)
            st.paragraph_format.space_after = Pt(0)

    for name in ("Front Matter Heading", "FrontMatterHeading"):
        if name in [s.name for s in doc.styles]:
            st = doc.styles[name]
            st.font.name = TNR
            st.font.size = Pt(14)
            st.font.bold = True
            st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            st.paragraph_format.line_spacing = 1.0


def enforce_fonts(doc: Document) -> None:
    for p in doc.paragraphs:
        for run in p.runs:
            set_run_font(run)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        set_run_font(run)
    for sec in doc.sections:
        for part in (sec.header, sec.footer):
            for p in part.paragraphs:
                for run in p.runs:
                    set_run_font(run)
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                set_run_font(run)


def configure_main_paragraphs(doc: Document) -> None:
    in_main = False
    in_refs = False
    for p in doc.paragraphs:
        style_name = p.style.name if p.style else ""
        if style_name in ("TOC 1", "TOC1", "TOC 2", "TOC2", "TOC 3", "TOC3", "Table of Figures", "TableofFigures"):
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.keep_with_next = False
            p.paragraph_format.keep_together = False
        if p.style and p.style.name == "Heading 1" and p.text.strip() == "Introduction":
            in_main = True
        if not in_main:
            continue
        if p.style and p.style.name == "Heading 1" and p.text.strip() == "References":
            in_refs = True
        if style_name in ("Normal", "Body Text", ""):
            p.paragraph_format.line_spacing = 1.5 if not in_refs else 1.0
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if not in_refs else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.widow_control = True
        if style_name.startswith("Heading"):
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.keep_together = True
        if style_name in ("Caption", "Caption 2", "Caption2"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.keep_together = True


def configure_page_numbering(doc: Document) -> None:
    sections = list(doc.sections)
    for idx, sec in enumerate(sections):
        sectpr = sec._sectPr
        pgnum = sectpr.find(qn("w:pgNumType"))
        if idx == 2:
            if pgnum is None:
                pgnum = OxmlElement("w:pgNumType")
                sectpr.append(pgnum)
            pgnum.set(qn("w:start"), "1")
            pgnum.set(qn("w:fmt"), "decimal")
        elif idx > 2 and pgnum is not None:
            pgnum.attrib.pop(qn("w:start"), None)


def enable_field_updates(doc: Document) -> None:
    settings = doc.settings._element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        existing = OxmlElement("w:updateFields")
        settings.append(existing)
    existing.set(qn("w:val"), "true")


def replace_theme_fonts(path: Path) -> None:
    tmp = path.with_suffix(".themefix.docx")
    with zipfile.ZipFile(path, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.startswith("word/theme/") or item.filename == "word/styles.xml":
                text = data.decode("utf-8")
                for old in ("Aptos Display", "Aptos", "Calibri Light", "Calibri", "Cambria"):
                    text = text.replace(old, TNR)
                data = text.encode("utf-8")
            zout.writestr(item, data)
    tmp.replace(path)


def build(source_path: Path, figures_dir: Path, output_path: Path) -> None:
    doc = Document(str(source_path))
    body = doc.element.body
    blocks = [copy.deepcopy(ch) for ch in body]
    final_sectpr = copy.deepcopy(body.sectPr)
    for child in list(body):
        body.remove(child)
    body.append(final_sectpr)

    def append(
        index: int,
        text: str | None = None,
        image_width: float | None = None,
        remove_page_break_before: bool = False,
    ):
        element = clone_with_plain_text(blocks[index], text) if text is not None else copy.deepcopy(blocks[index])
        if image_width is not None:
            resize_inline_drawing(element, image_width)
        if remove_page_break_before:
            ppr = element.find(qn("w:pPr"))
            if ppr is not None:
                page_break = ppr.find(qn("w:pageBreakBefore"))
                if page_break is not None:
                    ppr.remove(page_break)
        add_before_sectpr(doc, element)

    # Cover, rubric and all automatic front matter, ending with the section break that starts the main report.
    for i in range(0, 91):
        if i != 68:  # combine List of Figures and List of Tables on one front-matter page
            append(i)

    # Introduction.
    for i in range(91, 103):
        append(i)
    append(103, "The primary user is a student who can register, authenticate, recover a password and manage only records associated with the session user_id. An administrator authenticates through the same users table and can view registered accounts and aggregate module summaries. Administration is intentionally narrow because the assignment emphasizes the four student modules; it remains documented wherever role-based navigation, database design, requirements and implementation are relevant.")

    # Site hierarchy prose remains portrait; the hierarchy becomes a dedicated landscape plate.
    for i in (104, 105, 108, 109, 110, 111, 112, 113, 114):
        append(i)
    new_section(doc, True)
    add_picture(doc, figures_dir / "figure-1.png")
    append(107)

    # Flowchart introduction portrait, then four landscape plates.
    new_section(doc, False)
    for i in (115, 116, 117, 118):
        append(i)
    new_section(doc, True)
    for fig_no, caption_idx in ((2, 121), (3, 126), (4, 131), (5, 136)):
        add_picture(doc, figures_dir / f"figure-{fig_no}.png", page_break_before=fig_no != 2)
        append(caption_idx)

    # Step explanations are consolidated on two portrait pages.
    new_section(doc, False)
    for i in (119, 122, 123, 124, 127, 128):
        append(i)
    doc.add_page_break()
    for i in (129, 132, 133, 134, 137, 138):
        append(i)

    # Database introduction portrait, followed by ERD and two schema plates.
    for i in (139, 140):
        append(i)
    append(144, remove_page_break_before=True)
    append(143)
    new_section(doc, True)
    for fig_no, caption_idx in ((6, 142), (7, 146), (8, 148)):
        add_picture(doc, figures_dir / f"figure-{fig_no}.png", page_break_before=fig_no != 6)
        append(caption_idx)
    new_section(doc, False)
    for i in (149, 150, 151, 152, 153, 154):
        append(i)

    # Functional requirements remain unchanged.
    for i in range(155, 181):
        append(i)

    # Implementation prose remains portrait; architecture becomes a dedicated landscape plate.
    for i in (181, 182, 185, 186, 187, 188, 189):
        append(i)
    new_section(doc, True)
    add_picture(doc, figures_dir / "figure-9.png")
    append(184)
    new_section(doc, False)
    for i in range(190, 196):
        append(i)
    append(196, image_width=3.5)
    append(197)
    append(198, image_width=6.0)
    append(199)
    append(200, image_width=3.5)
    append(201)

    # The standalone Testing and Debugging chapter is intentionally omitted.
    for i in (216, 217):
        append(i)
    append(218, "The implementation also applies session timeout, safe cookie attributes, strong password hashing and recovery, CSRF validation, output encoding, prepared statements, ownership constraints and controlled exception logging. Optional filters, exports, drafts, savings goals and progress analysis improve usability without replacing the mandatory functions.")
    append(219)
    for i in range(220, 231):
        append(i)

    configure_styles(doc)
    for idx, table in enumerate(doc.tables):
        set_table_rules(table, report_table=idx >= 2)
    configure_main_paragraphs(doc)
    enforce_fonts(doc)
    enable_field_updates(doc)
    configure_page_numbering(doc)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    replace_theme_fonts(output_path)
    print(f"Built report with {len(doc.sections)} sections, {len(doc.tables)} tables and {len(doc.inline_shapes)} inline figures.")


def main() -> None:
    if len(sys.argv) != 4:
        raise SystemExit("Usage: build_report.py SOURCE_DOCX WORD_FIGURE_DIR OUTPUT_DOCX")
    build(Path(sys.argv[1]), Path(sys.argv[2]), Path(sys.argv[3]))


if __name__ == "__main__":
    main()
